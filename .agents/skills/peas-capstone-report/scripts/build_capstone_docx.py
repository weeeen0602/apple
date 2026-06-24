#!/usr/bin/env python3
"""Build report/專題報告.docx from MD with embedded PNG images."""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
except ImportError:
    print("ERROR: python-docx required. Run: uv add python-docx", file=sys.stderr)
    raise SystemExit(1)


IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
PAGE_WIDTH_CM = 15.0


def parse_md(md_text: str) -> dict:
    lines = md_text.splitlines()
    blocks: list[tuple[str, str, list]] = []
    title = "專題報告"
    meta_lines: list[str] = []
    current_level = 0
    current_title = ""
    current_body: list = []

    def flush() -> None:
        nonlocal current_title, current_body, current_level
        if current_title:
            blocks.append((current_level, current_title, current_body.copy()))
        current_title = ""
        current_body = []
        current_level = 0

    for line in lines:
        if line.startswith("# ") and not blocks and not current_title:
            title = line[2:].strip()
            continue
        if line.startswith("**組別") or line.startswith("**日期"):
            meta_lines.append(line.strip("* "))
            continue
        if line.strip() == "---":
            continue
        m = HEADING_RE.match(line)
        if m:
            flush()
            current_level = len(m.group(1))
            current_title = m.group(2).strip()
            continue
        img = IMAGE_RE.match(line.strip())
        if img:
            current_body.append(("image", img.group(1), img.group(2).strip()))
            continue
        b = BULLET_RE.match(line)
        if b:
            current_body.append(("bullet", b.group(1).strip()))
            continue
        if line.strip() and current_title:
            current_body.append(("para", line.strip()))

    flush()
    return {"title": title, "meta": meta_lines, "blocks": blocks}


def add_image(doc: Document, report_dir: Path, rel_path: str, caption: str = "") -> bool:
    img_path = (report_dir / rel_path).resolve()
    if not img_path.is_file():
        print(f"WARN: missing image {img_path}", file=sys.stderr)
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(PAGE_WIDTH_CM * 0.8))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(10)
    return True


def build_docx(report_dir: Path) -> tuple[Path, int]:
    md_path = report_dir / "專題報告.md"
    if not md_path.is_file():
        raise FileNotFoundError(f"Missing {md_path}")

    data = parse_md(md_path.read_text(encoding="utf-8"))
    doc = Document()

    h = doc.add_heading(data["title"], level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for meta in data["meta"]:
        p = doc.add_paragraph(meta)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    image_count = 0
    for level, heading, body in data["blocks"]:
        doc.add_heading(heading, level=min(level, 3))
        for item in body:
            if item[0] == "bullet":
                doc.add_paragraph(item[1], style="List Bullet")
            elif item[0] == "para":
                doc.add_paragraph(item[1])
            elif item[0] == "image":
                _, alt, rel = item
                if add_image(doc, report_dir, rel, alt):
                    image_count += 1

    out = report_dir / "專題報告.docx"
    doc.save(str(out))
    return out, image_count


def count_docx_images(docx_path: Path) -> int:
    """Count inline image embeds in document body."""
    with zipfile.ZipFile(docx_path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return xml.count("r:embed")


def verify_docx(report_dir: Path, min_demo: int = 0) -> int:
    docx_path = report_dir / "專題報告.docx"
    if not docx_path.is_file():
        print(f"ERROR: missing {docx_path}", file=sys.stderr)
        return 1

    assets = report_dir / "assets"
    demo_count = len(list(assets.glob("demo-*.png")))
    required = 2 + max(demo_count, min_demo)
    embedded = count_docx_images(docx_path)

    if embedded < required:
        print(
            f"ERROR: docx has {embedded} images, need at least {required} "
            f"(server + architecture + {demo_count} demo)",
            file=sys.stderr,
        )
        return 1

    print(f"OK: verify docx images={embedded} (required>={required})")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--report-dir", type=Path, default=Path("report"))
    parser.add_argument(
        "--verify-only",
        action="store_true",
        help="Only verify existing docx embedded images",
    )
    parser.add_argument(
        "--min-demo",
        type=int,
        default=0,
        help="Minimum demo count for verify (from preflight)",
    )
    args = parser.parse_args()

    if args.verify_only:
        return verify_docx(args.report_dir, args.min_demo)

    try:
        out, image_count = build_docx(args.report_dir)
    except FileNotFoundError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    embedded = count_docx_images(out)
    demo_count = len(list((args.report_dir / "assets").glob("demo-*.png")))
    required = 2 + demo_count
    if embedded < required:
        print(
            f"ERROR: built docx has {embedded} images, need at least {required}",
            file=sys.stderr,
        )
        return 1

    print(f"OK: {out} (images={embedded}, parsed={image_count})")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
